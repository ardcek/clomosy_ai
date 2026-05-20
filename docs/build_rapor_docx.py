#!/usr/bin/env python3
"""Teknik rapor Markdown yerine dogrudan Word (.docx) uretir. Calistir: python3 build_rapor_docx.py"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_paragraph_spacing(paragraph, before_pt=0, after_pt=6, line_multiple=1.15):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    pf.line_spacing = line_multiple


def shade_table_header(row):
    for cell in row.cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "EFF6FF")
        tc_pr.append(shd)


def main():
    out = Path(__file__).resolve().parent / "proje-teknik-raporu-akilli-ariza.docx"
    doc = Document()

    sec = doc.sections[0]
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.85)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "PROJE TEKNİK RAPORU\n"
        "AKILLI ARIZA & SAHA OPERASYON\n"
        "(Clomosy / TRObject)"
    )
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(15, 23, 42)
    set_paragraph_spacing(title, before_pt=0, after_pt=14)

    meta = [
        ("Hazırlayan", "Arda Çekiç, Esra Ağlar"),
        ("E-posta", "ardcek@hotmail.com, llesraaq.nur@gmail.com"),
        ("Platform", "Clomosy / TRObject"),
        (
            "Veri & kimlik",
            "Google Firebase (Identity Toolkit + Realtime Database), REST tabanlı erişim",
        ),
        ("Proje aktivasyon kodu", "EDUCA-7F7DC"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(meta):
        row = table.rows[i]
        row.cells[0].text = k
        row.cells[1].text = v
        for c in (0, 1):
            p = row.cells[c].paragraphs[0]
            for r in p.runs:
                r.font.name = "Calibri"
                r.font.size = Pt(10.5)
                if c == 0:
                    r.bold = True
                    r.font.color.rgb = RGBColor(71, 85, 105)
                else:
                    r.font.color.rgb = RGBColor(15, 23, 42)
    shade_table_header(table.rows[0])
    doc.add_paragraph()

    # --- Demo giris hesaplari (rapor / degerlendirme) ---
    gh = doc.add_paragraph()
    gr = gh.add_run("GİRİŞ BİLGİLERİ (DEMO / TEST)")
    gr.bold = True
    gr.font.size = Pt(13)
    gr.font.color.rgb = RGBColor(37, 99, 235)
    set_paragraph_spacing(gh, before_pt=4, after_pt=6)

    gp = doc.add_paragraph(
        "Asagidaki hesaplar rapor ve demo amaclidir; canli ortamda kullanilmamalidir."
    )
    set_paragraph_spacing(gp, before_pt=0, after_pt=10, line_multiple=1.2)
    for run in gp.runs:
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(100, 116, 139)

    login_rows = [
        ("Rol", "E-posta", "Sifre"),
        ("Yonetici", "yonetici@firma.com", "yonetici123"),
        ("Talep acan", "talepacan@firma.com", "talep123"),
        ("Saha personeli", "saha@firma.com", "saha123"),
    ]
    lt = doc.add_table(rows=len(login_rows), cols=3)
    lt.style = "Table Grid"
    lt.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row_data in enumerate(login_rows):
        row = lt.rows[ri]
        for ci, val in enumerate(row_data):
            row.cells[ci].text = val
            p = row.cells[ci].paragraphs[0]
            for r in p.runs:
                r.font.name = "Calibri"
                r.font.size = Pt(10.5)
                if ri == 0:
                    r.bold = True
                    r.font.color.rgb = RGBColor(71, 85, 105)
                elif ci == 2:
                    r.font.name = "Consolas"
                    r.font.color.rgb = RGBColor(15, 23, 42)
                else:
                    r.font.color.rgb = RGBColor(15, 23, 42)
    shade_table_header(lt.rows[0])
    doc.add_paragraph()

    sections = [
        (
            "1. YÖNETİCİ ÖZETİ",
            [
                "Akıllı Arıza & Saha Operasyon, Clomosy ortamında TRObject ile geliştirilen, "
                "sahada ve ofiste arıza / iş taleplerinin dijitalleştirilmesini hedefleyen düşük kodlu bir çözümdür. "
                "Talep oluşturma, yönetici tarafından personele atama ve saha personelinin durum güncellemeleri tek "
                "uygulama akışında toplanır; kullanıcılar rol bazlı ekranlara yönlendirilir. Amaç: taleplerin merkezi "
                "ve güncel kalması, iletişim ve manuel takip yükünün azaltılması."
            ],
        ),
        (
            "2. TEKNİK ALTYAPI VE MİMARİ",
            [
                "Geliştirme ortamı: İş mantığı ve arayüz Clomosy üzerinde TRObject (.tro) birimleriyle tanımlanır; "
                "formlar, Pro bileşenler ve olay bağları platform kalıplarına uygundur.",
                "Kimlik ve yetkilendirme: Firebase Authentication (e-posta / şifre) ile oturum; kullanıcı profili ve "
                "rol bilgisi (Yonetici, TalepAcan, SahaPersoneli) Realtime Database’deki kullanıcı kaydı üzerinden "
                "okunur ve ilgili modüle yönlendirme yapılır.",
                "Veri yönetimi: Talep ve operasyon verileri Firebase Realtime Database üzerinde tutulur; istemci "
                "tarafında REST çağrıları ile okuma, yazma ve sorgu (ör. orderBy + equalTo) gerçekleştirilir. "
                "Mobil ve masaüstü istemciler aynı canlı veri modeline erişebilir.",
                "Arayüz: ProPanel, ProListView, ProEdit vb. bileşenlerle kart tabanlı, responsive yerleşim; rol "
                "ekranlarında ortak tasarım token’ları (renk, tipografi) kullanılır.",
            ],
        ),
        (
            "3. MODÜLER YAPI VE FONKSİYONEL ÖZELLİKLER",
            [
                "3.1. Güvenlik ve giriş (uLogin): E-posta / şifre ile giriş; hata mesajlarının kullanıcı dostu "
                "gösterimi; başarılı girişte role göre uYoneticiAtama, uTalepAcma veya uSahaPersoneli birimine geçiş.",
                "3.2. Talep açan modül (uTalepAcma): Yeni talep formu (başlık, açıklama, konum, öncelik); talebin "
                "RTDB’ye yazılması; kullanıcıya ait taleplerin listelenmesi, yenileme ve metin araması ile filtreleme. "
                "Önceliğe göre istemci tarafında SLA süresi hesaplanır.",
                "3.3. Yönetici atama modülü (uYoneticiAtama): Açık taleplerin listelenmesi; personel seçimi ve "
                "talebe atama; talep özeti ve arama.",
                "3.4. Saha personeli modülü (uSahaPersoneli): Personele atanan taleplerin listelenmesi; iş kuralına "
                "uygun durum geçişleri (ör. Atandı → Yolda / İşlemde, İşlemde → Beklemede / Tamamlandı vb.); "
                "isteğe bağlı not ve durum güncellemesi için RTDB güncellemeleri.",
            ],
        ),
        (
            "4. PROJE ÇIKTILARI VE AVANTAJLAR",
            None,
        ),
        (
            "5. SONUÇ",
            [
                "Akıllı Arıza & Saha Operasyon, Clomosy ve TRObject ile geliştirilen, talep yaşam döngüsünü "
                "açma → atama → saha durumu olarak kapsayan uygulamadır. Firebase ile kimlik ve gerçek zamanlı veri "
                "birleştirilerek, küçük ve orta ölçekli operasyon ekipleri için uygulanabilir bir saha operasyon "
                "iskeleti sunulmaktadır."
            ],
        ),
    ]

    bullets = [
        "Hız: TRObject ile ekran ve iş akışlarının tek kod tabanında hızlı iterasyonu; Clomosy bileşen modeli ile "
        "form yoğun geliştirmenin kısaltılması.",
        "Erişilebilirlik: Firebase tabanlı veri; sahada mobil, ofiste aynı mantıkla kullanım.",
        "Maliyet / bakım: Ağır özel API zorunluluğu olmadan RTDB ve kimlik servisleriyle sade mimari; modüllerin "
        "rol bazlı ayrılması bakımı kolaylaştırır.",
    ]

    for heading, paras in sections:
        h = doc.add_paragraph()
        r = h.add_run(heading)
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(37, 99, 235)
        set_paragraph_spacing(h, before_pt=14, after_pt=8)

        if paras:
            for text in paras:
                p = doc.add_paragraph(text)
                p.paragraph_format.first_line_indent = Pt(0)
                set_paragraph_spacing(p, before_pt=0, after_pt=8, line_multiple=1.2)
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(11)
        elif heading.startswith("4."):
            for text in bullets:
                p = doc.add_paragraph(text, style="List Bullet")
                set_paragraph_spacing(p, before_pt=0, after_pt=6, line_multiple=1.2)
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(11)

    foot = doc.add_paragraph()
    set_paragraph_spacing(foot, before_pt=18, after_pt=4)
    r1 = foot.add_run("İletişim: ")
    r1.bold = True
    r1.font.color.rgb = RGBColor(71, 85, 105)
    r2 = foot.add_run("ardcek@hotmail.com  ·  llesraaq.nur@gmail.com")
    r2.font.color.rgb = RGBColor(51, 65, 85)

    foot2 = doc.add_paragraph()
    r3 = foot2.add_run("Proje aktivasyon kodu: ")
    r3.bold = True
    r3.font.color.rgb = RGBColor(71, 85, 105)
    r4 = foot2.add_run("EDUCA-7F7DC")
    r4.font.name = "Consolas"
    r4.font.size = Pt(11)
    r4.font.color.rgb = RGBColor(15, 23, 42)

    doc.save(out)
    print(f"Yazildi: {out}")


if __name__ == "__main__":
    main()
